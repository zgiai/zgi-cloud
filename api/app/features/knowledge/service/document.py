import os
import hashlib
from typing import List, Optional, Dict, Any, BinaryIO, Tuple
from fastapi import UploadFile
from sqlalchemy import select
from sqlalchemy.orm import Session
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.features.knowledge.models.document import Document, DocumentStatus, DocumentChunk
from app.features.knowledge.models.knowledge import KnowledgeBase
from app.features.knowledge.schemas.request.document import DocumentUpdate
from app.features.knowledge.service.knowledge import KnowledgeBaseService
from app.features.knowledge.core.config import get_document_settings
from app.features.knowledge.core.response import (
    ServiceResponse,
    ValidationError,
    NotFoundError, ServiceError
)
from app.features.knowledge.core.decorators import (
    handle_service_errors,
    retry
)
from app.features.knowledge.core.logging import (
    service_logger,
    audit_logger,
    metrics_logger
)

class DocumentService:
    """Service for managing documents in knowledge bases"""

    def __init__(
        self,
        db: Session,
        kb_service: KnowledgeBaseService
    ):
        self.db = db
        self.kb_service = kb_service
        self.settings = get_document_settings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.CHUNK_SIZE,
            chunk_overlap=self.settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    @retry(max_attempts=3)
    async def _extract_text(self, file: BinaryIO, file_type: str) -> str:
        """Extract text from a file"""
        try:
            if file_type == "pdf":
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n\n"
                return text
            
            elif file_type == "docx":
                doc = docx.Document(file)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n\n"
                return text
            
            elif file_type == "txt":
                return file.read().decode("utf-8")
            
            else:
                raise ValidationError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            service_logger.error(
                f"Text extraction failed for file type {file_type}",
                "text_extraction_failed",
                error=str(e)
            )
            raise

    async def _split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        chunks = self.text_splitter.split_text(text)
        if len(chunks) > self.settings.MAX_CHUNKS_PER_DOC:
            raise ValidationError(
                f"Document too large. Maximum chunks allowed: {self.settings.MAX_CHUNKS_PER_DOC}"
            )
        return chunks

    @retry(max_attempts=3)
    async def _store_file(
        self,
        file: UploadFile,
        kb_id: int,
        user_id: int
    ) -> tuple[str, str]:
        """Store uploaded file"""
        try:
            # Validate file size
            file_content = await file.read()
            if len(file_content) > self.settings.MAX_FILE_SIZE:
                raise ValidationError(
                    f"File too large. Maximum size allowed: {self.settings.MAX_FILE_SIZE} bytes"
                )
            
            # Create directory if not exists
            upload_dir = os.path.join(
                self.settings.UPLOAD_DIR,
                str(user_id),
                str(kb_id)
            )
            os.makedirs(upload_dir, exist_ok=True)
            
            # Generate unique filename
            file_hash = hashlib.md5(file_content).hexdigest()
            file_ext = os.path.splitext(file.filename)[1]
            filename = f"{file_hash}{file_ext}"
            filepath = os.path.join(upload_dir, filename)
            
            # Write file
            with open(filepath, "wb") as f:
                f.write(file_content)
            
            return filepath, file_hash
            
        except Exception as e:
            service_logger.error(
                "File storage failed",
                "file_storage_failed",
                error=str(e),
                kb_id=kb_id,
                user_id=user_id
            )
            raise

    @handle_service_errors
    async def upload_document(
        self,
        kb_id: int,
        file: UploadFile,
        user_id: int,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ServiceResponse[Document]:
        """Upload and process a document"""
        start_time = metrics_logger.time()
        
        try:
            # Validate file type
            file_ext = os.path.splitext(file.filename)[1].lower()
            if file_ext[1:] not in self.settings.SUPPORTED_TYPES:
                raise ValidationError(
                    f"Unsupported file type. Supported types: {', '.join(self.settings.SUPPORTED_TYPES)}"
                )
            
            # Get knowledge base
            # kb_response = await self.kb_service.get_knowledge_base(kb_id, user_id)
            # if not kb_response.success:
            #     raise NotFoundError(
            #         f"Knowledge base not found",
            #         details={"kb_id": kb_id}
            #     )
            # kb = KnowledgeBase(**dict(kb_response.data))
            kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
            if not kb:
                raise NotFoundError("Knowledge base not found")
            
            try:
                # Store file
                filepath, file_hash = await self._store_file(file, kb_id, user_id)
                # Create document record with processing status
                document = Document(
                    kb_id=kb_id,
                    file_name=file.filename,
                    file_path=filepath,
                    file_type=file_ext[1:],
                    file_size=file.size,
                    file_hash=file_hash,
                    status=DocumentStatus.PROCESSING.value,
                    meta_info=metadata
                )
                self.db.add(document)
                self.db.commit()
                self.db.refresh(document)
                
                # Extract text
                with open(filepath, "rb") as f:
                    text = await self._extract_text(f, file_ext[1:])
                
                # Split text into chunks
                chunks = await self._split_text(text)
                
                # Store chunks in the database
                for i, chunk in enumerate(chunks):
                    document_chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=i,
                        content=chunk,
                        token_count=len(chunk.split())
                    )
                    self.db.add(document_chunk)
                
                self.db.commit()
                
                # Generate embeddings and store vectors
                embeddings = await self.kb_service.embedding.get_embeddings(chunks)
                chunk_metadata = [{
                    "document_id": document.id,
                    "chunk_index": i,
                    "text": chunk,
                    **(metadata or {})
                } for i, chunk in enumerate(chunks)]
                
                success = await self.kb_service.vector_db.insert_vectors(
                    collection_name=kb.collection_name,
                    vectors=embeddings,
                    metadata=chunk_metadata
                )
                
                if not success:
                    raise Exception("Failed to store vectors")
                
                # Update document status
                document.status = DocumentStatus.COMPLETED.value
                document.chunk_count = len(chunks)
                self.db.commit()
                
                # Update knowledge base statistics
                kb.document_count += 1
                kb.total_chunks += len(chunks)
                self.db.commit()
                
                audit_logger.log_access(
                    user_id,
                    "document",
                    str(document.id),
                    "upload",
                    "success",
                    details={"kb_id": kb_id}
                )
                
                metrics_logger.log_operation(
                    "document_upload",
                    metrics_logger.time() - start_time,
                    True,
                    {"document_id": document.id, "chunks": len(chunks)}
                )
                
                return ServiceResponse.ok(document.to_dict())
                
            except Exception as e:
                # Update document status on failure
                document.status = DocumentStatus.FAILED.value
                document.error_message = str(e)
                self.db.commit()
                
                # Clean up file if stored
                if document.file_path and os.path.exists(document.file_path):
                    try:
                        os.remove(document.file_path)
                    except:
                        pass
                
                metrics_logger.log_operation(
                    "document_upload",
                    metrics_logger.time() - start_time,
                    False,
                    {"error": str(e)}
                )
                raise
                
        except Exception as e:
            self.db.rollback()
            raise

    # @handle_service_errors
    async def get_document(
        self,
        doc_id: int,
        user_id: int
    ) -> Document:
        """Get a document by ID"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")
            
        # Check access through knowledge base
        kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        if not kb_response.success:
            raise ServiceError(f"Knowledge base {document.kb_id} not found")
            
        audit_logger.log_access(
            user_id,
            "document",
            str(doc_id),
            "read",
            "success"
        )
        
        return document

    @handle_service_errors
    async def delete_document(
        self,
        doc_id: int,
        user_id: int
    ) -> ServiceResponse[None]:
        """Delete a document"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")
            
        # Check access through knowledge base
        # kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        # if not kb_response.success:
        #     return kb_response
        # kb = kb_response.data
        kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == document.kb_id).first()
        if not kb:
            raise NotFoundError("Knowledge base not found")
        
        try:
            # Delete vectors from vector database
            success = await self.kb_service.vector_db.delete_vectors(
                collection_name=kb.collection_name,
                vector_ids=[]
                # filter={"document_id": doc_id}
            )
            
            if not success:
                raise Exception("Failed to delete vectors")
            
            # Delete file
            if document.file_path and os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Update knowledge base statistics
            kb.document_count -= 1
            kb.total_chunks -= document.chunk_count
            
            # Delete document record
            self.db.delete(document)
            self.db.commit()
            
            audit_logger.log_access(
                user_id,
                "document",
                str(doc_id),
                "delete",
                "success"
            )
            
            return ServiceResponse.ok()
            
        except Exception as e:
            self.db.rollback()
            raise

    # @handle_service_errors
    async def list_documents(
        self,
        kb_id: int,
        user_id: int,
        skip: int = 0,
        limit: int = 10,
        file_type: Optional[str] = None,
        status: Optional[int] = None,
        search: Optional[str] = None
    ) -> Tuple[List[Document], int]:
        """List documents in a knowledge base"""
        # Check access through knowledge base
        # kb_response = await self.kb_service.get_knowledge_base(kb_id, user_id)
        # if not kb_response.success:
        #     return kb_response
        kb = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == kb_id).first()
        if not kb:
            raise NotFoundError("Knowledge base not found")
            
        # Build query
        query = self.db.query(Document).filter(Document.kb_id == kb_id)
        
        if file_type:
            query = query.filter(Document.file_type == file_type)
        if status is not None:
            query = query.filter(Document.status == status)

        if search:
            query = query.filter(Document.file_name.ilike(f"%{search}%"))

        total = query.count()
        documents = query.offset(skip).limit(limit).all()
        
        return documents, total

    async def update_document(
        self,
        doc_id: int,
        update_data: DocumentUpdate,
        user_id: int
    ):
        """Update a document"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")

        kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        if not kb_response.success:
            raise ServiceError(f"Knowledge base {document.kb_id} not found")
        
        for field, value in update_data.dict(exclude_unset=True).items():
            setattr(document, field, value)
        
        # 提交更改
        self.db.commit()
        self.db.refresh(document)
        
        # 记录审计日志
        audit_logger.log_access(
            user_id,
            "document",
            str(doc_id),
            "update",
            "success"
        )
        
        return document

    # @handle_service_errors
    async def list_chunks(
        self,
        doc_id: int,
        user_id: int,
        skip: int = 0,
        limit: int = 10
    ) -> Tuple[List[DocumentChunk], int]:
        """List chunks of a document"""
        document = self.db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise NotFoundError(f"Document {doc_id} not found")
        
        # Check access through knowledge base
        kb_response = await self.kb_service.get_knowledge_base(document.kb_id, user_id)
        if not kb_response.success:
            raise ServiceError(f"Knowledge base {document.kb_id} not found")
        
        query = self.db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id)
        total = query.count()
        chunks = query.offset(skip).limit(limit).all()
        
        return chunks, total
